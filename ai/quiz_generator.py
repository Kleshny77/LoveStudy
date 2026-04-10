# Генерация викторин по файлам/материалам через LLM.
#
# Приоритет: Gemini-compatible endpoint, затем DeepSeek, затем Groq (Llama), затем Gemini.
# Поддерживаемые форматы: PDF, DOCX, TXT, изображения (Gemini vision).
# Вопросы нетривиальные: все варианты ответов правдоподобны, только один верный.

from __future__ import annotations

import io
import json
import logging
import random
import re
from dataclasses import dataclass

import httpx

logger = logging.getLogger(__name__)

# Cloudflare часто режет python-httpx / Python-urllib на датацентровых IP.
_GEMINI_PROXY_HEADERS_EXTRA = {
    "User-Agent": "Mozilla/5.0 (compatible; LoveStudy/1.0; +https://github.com/Kleshny77/LoveStudy)",
    "Accept": "application/json",
}

# Лимиты Telegram для quiz
MAX_QUESTION_LEN = 300
MAX_OPTION_LEN = 72
MAX_OPTIONS = 5
MIN_OPTIONS = 3
MAX_EXPLANATION_LEN = 200


@dataclass(slots=True)
class QuizQuestion:
    question: str
    options: list[str]
    correct_index: int
    explanation: str


def _truncate(s: str, limit: int) -> str:
    s = (s or "").strip()
    if len(s) <= limit:
        return s
    return s[: limit - 1].rstrip() + "…"


def _normalize_question_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _extract_text_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n\n".join(parts).strip() or ""
    except Exception as e:
        logger.warning("pypdf extract failed: %s", e)
        return ""


def _extract_text_docx(data: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts).strip() or ""
    except Exception as e:
        logger.warning("python-docx extract failed: %s", e)
        return ""


def _extract_text_txt(data: bytes) -> str:
    try:
        return data.decode("utf-8", errors="replace").strip()
    except Exception:
        return ""


def _fetch_url_text(url: str, max_chars: int = 50000) -> str:
    try:
        with httpx.Client(timeout=15.0, follow_redirects=True) as client:
            r = client.get(url)
            r.raise_for_status()
            html = r.text
    except Exception as e:
        logger.warning("fetch url %s failed: %s", url[:80], e)
        return ""

    # Простая очистка HTML
    text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:max_chars] if text else ""


def extract_text_from_material(
    data: bytes | None,
    url: str | None,
    mime_type: str | None,
    filename: str | None,
) -> str:
    """Извлекает текст из материала для передачи в LLM."""
    if url:
        return ""

    if not data:
        return ""

    mime = (mime_type or "").lower()
    name = (filename or "").lower()

    if "pdf" in mime or name.endswith(".pdf"):
        return _extract_text_pdf(data)
    if any(
        x in mime or name.endswith(x)
        for x in (".docx", ".doc", "word", "document")
    ):
        return _extract_text_docx(data)
    if "text" in mime or name.endswith((".txt", ".md", ".rst")):
        return _extract_text_txt(data)

    # Пробуем по расширению при неизвестном MIME
    if "octet-stream" in mime or not mime:
        if name.endswith(".pdf"):
            return _extract_text_pdf(data)
        if name.endswith(".docx"):
            return _extract_text_docx(data)
        if name.endswith(".txt"):
            return _extract_text_txt(data)

    return _extract_text_txt(data)


_QUIZ_SYSTEM = """Ты — эксперт по созданию образовательных викторин. На основе контента создай вопросы для проверки понимания.

КРИТИЧЕСКИ ВАЖНО:
1. Все варианты ответов ПРАВДОПОДОБНЫ и связаны с темой. Запрещены очевидно неверные (типа "банан" на вопрос про код).
2. Неправильные ответы — такие, что человек мог бы выбрать по недопониманию: похожие термины, частично верные утверждения, типичные ошибки.
3. Вопросы проверяют понимание ключевых идей, не тривиальные факты.
4. Объяснение — кратко (до 200 символов), почему правильный ответ верен. На русском.
5. Варианты ответов формулируй коротко и ясно, желательно до 72 символов каждый.

Верни ТОЛЬКО валидный JSON без markdown-обёртки:
{
  "questions": [
    {
      "question": "текст вопроса",
      "options": ["вариант A", "вариант B", "вариант C", "вариант D"],
      "correct_index": 0,
      "explanation": "краткое объяснение"
    }
  ]
}
Вариантов: 3–5. correct_index — индекс правильного (0-based)."""


def _build_quiz_prompt(num_questions: int, exclude_questions: list[str] | None = None) -> str:
    prompt = f"{_QUIZ_SYSTEM}\n\nСоздай ровно {num_questions} вопросов."
    if exclude_questions:
        limited = [q.strip() for q in exclude_questions if q.strip()][:20]
        if limited:
            prompt += "\n\nКРИТИЧЕСКИ ВАЖНО: не повторяй следующие вопросы и их формулировки:\n"
            prompt += "\n".join(f"- {question}" for question in limited)
            prompt += "\nПридумай новые вопросы по другим деталям и аспектам материала."
    return prompt


def _call_deepseek(
    api_key: str,
    content_text: str,
    num_questions: int = 3,
    exclude_questions: list[str] | None = None,
) -> str:
    """Генерация викторины через DeepSeek. Только текст."""
    try:
        user_msg = f"Создай ровно {num_questions} вопросов по следующему контенту.\n\n{content_text[:110000]}"
        response = httpx.post(
            "https://api.deepseek.com/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": _build_quiz_prompt(num_questions, exclude_questions)},
                    {"role": "user", "content": user_msg},
                ],
                "temperature": 0.4,
                "max_tokens": 4096,
            },
            timeout=60.0,
        )
        response.raise_for_status()
        data = response.json()
        text = ((data.get("choices") or [{}])[0].get("message") or {}).get("content", "").strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        return text or "{}"
    except Exception as e:
        logger.warning("DeepSeek generate failed: %s", e)
        return "{}"


def _call_gemini_proxy(
    api_key: str,
    base_url: str,
    model: str,
    content_text: str,
    num_questions: int = 3,
    exclude_questions: list[str] | None = None,
) -> str:
    """Генерация викторины через Gemini-compatible endpoint (OpenAI-style /chat/completions).

    Часть прокси (в т.ч. под Gemini) не принимает role=system — шлём один user с инструкцией и контентом.
    """
    try:
        user_msg = f"Создай ровно {num_questions} вопросов по следующему контенту.\n\n{content_text[:110000]}"
        instructions = _build_quiz_prompt(num_questions, exclude_questions)
        combined_user = f"{instructions}\n\n{user_msg}"
        url = f"{base_url.rstrip('/')}/chat/completions"
        payload = {
            "model": model,
            "messages": [{"role": "user", "content": combined_user}],
            "temperature": 0.4,
            "max_tokens": 4096,
        }
        response = httpx.post(
            url,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                **_GEMINI_PROXY_HEADERS_EXTRA,
            },
            json=payload,
            timeout=120.0,
        )
        if response.status_code >= 400:
            snippet = (response.text or "")[:1500].replace("\n", " ")
            logger.warning(
                "Gemini proxy HTTP %s: %s",
                response.status_code,
                snippet or "(пустое тело)",
            )
        response.raise_for_status()
        data = response.json()
        text = ((data.get("choices") or [{}])[0].get("message") or {}).get("content", "").strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        return text or "{}"
    except Exception as e:
        logger.warning("Gemini proxy generate failed: %s", e)
        return "{}"


def _call_groq(
    api_key: str,
    content_text: str,
    num_questions: int = 3,
    exclude_questions: list[str] | None = None,
) -> str:
    """Генерация викторины через Groq (Llama). Только текст."""
    try:
        from groq import Groq

        client = Groq(api_key=api_key)
        user_msg = f"Создай ровно {num_questions} вопросов по следующему контенту.\n\n{content_text[:70000]}"
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": _build_quiz_prompt(num_questions, exclude_questions)},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.4,
            max_tokens=4096,
        )
        text = (response.choices[0].message.content or "").strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        return text or "{}"
    except Exception as e:
        logger.warning("Groq generate failed: %s", e)
        return "{}"


def _call_gemini(
    api_key: str,
    content_parts: list,
    num_questions: int = 3,
    exclude_questions: list[str] | None = None,
) -> str:
    import google.generativeai as genai
    from google.api_core.exceptions import ResourceExhausted

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-3.1-flash-lite-preview")
    prompt = _build_quiz_prompt(num_questions, exclude_questions)

    try:
        response = model.generate_content(content_parts + [prompt])
        if not response or not response.text:
            return "{}"
        text = response.text.strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        return text
    except ResourceExhausted as e:
        logger.warning("Gemini quota exceeded: %s", e)
        return "{}"
    except Exception as e:
        logger.exception("Gemini generate_content failed: %s", e)
        return "{}"


def _parse_quiz_json(raw: str, exclude_questions: list[str] | None = None) -> list[QuizQuestion]:
    questions: list[QuizQuestion] = []
    seen_questions = {_normalize_question_text(question) for question in (exclude_questions or []) if question}
    try:
        data = json.loads(raw)
        for q in data.get("questions", [])[:10]:
            opts = list(q.get("options") or [])
            if len(opts) < MIN_OPTIONS or len(opts) > MAX_OPTIONS:
                continue
            idx = int(q.get("correct_index", 0))
            if idx < 0 or idx >= len(opts):
                idx = 0
            expl = (q.get("explanation") or "").strip()
            quest = (q.get("question") or "").strip()
            if not quest:
                continue
            normalized_question = _normalize_question_text(quest)
            if not normalized_question or normalized_question in seen_questions:
                continue
            indexed_options = list(enumerate(opts))
            random.shuffle(indexed_options)
            shuffled_options = [option for _, option in indexed_options]
            shuffled_correct_index = next(
                (new_idx for new_idx, (old_idx, _) in enumerate(indexed_options) if old_idx == idx),
                0,
            )
            questions.append(
                QuizQuestion(
                    question=_truncate(quest, MAX_QUESTION_LEN),
                    options=[_truncate(o, MAX_OPTION_LEN) for o in shuffled_options],
                    correct_index=shuffled_correct_index,
                    explanation=_truncate(expl, MAX_EXPLANATION_LEN),
                )
            )
            seen_questions.add(normalized_question)
    except json.JSONDecodeError as e:
        logger.warning("Quiz JSON parse error: %s", e)
    return questions


def generate_quiz_from_text(
    text: str,
    subject_or_filename: str = "материал",
    num_questions: int = 3,
    exclude_questions: list[str] | None = None,
    *,
    gemini_proxy_key: str | None = None,
    gemini_proxy_base_url: str | None = None,
    gemini_proxy_model: str = "gemini-2.0-flash",
    deepseek_key: str | None = None,
    groq_key: str | None = None,
    gemini_key: str | None = None,
) -> list[QuizQuestion]:
    """Генерирует викторину по текстовому контенту. Сначала пробует Gemini-compatible endpoint."""
    if not text or len(text.strip()) < 100:
        return []

    content = f"Контент из «{subject_or_filename}»:\n\n{text[:80000]}"
    raw = "{}"
    requested_questions = min(8, num_questions + 3) if exclude_questions else num_questions
    if gemini_proxy_key and gemini_proxy_base_url:
        raw = _call_gemini_proxy(
            gemini_proxy_key,
            gemini_proxy_base_url,
            gemini_proxy_model,
            content,
            requested_questions,
            exclude_questions,
        )
    if (not raw or raw == "{}") and deepseek_key:
        raw = _call_deepseek(deepseek_key, content, requested_questions, exclude_questions)
    if (not raw or raw == "{}") and groq_key:
        raw = _call_groq(groq_key, content, requested_questions, exclude_questions)
    if (not raw or raw == "{}") and gemini_key:
        raw = _call_gemini(gemini_key, [content], requested_questions, exclude_questions)
    return _parse_quiz_json(raw, exclude_questions)[:num_questions]


def generate_quiz_from_image(
    image_bytes: bytes,
    mime_type: str,
    subject_or_filename: str = "изображение",
    num_questions: int = 2,
    exclude_questions: list[str] | None = None,
    *,
    gemini_key: str | None = None,
) -> list[QuizQuestion]:
    """Генерирует викторину по изображению через Gemini (Groq не поддерживает картинки)."""
    if not gemini_key:
        return []
    try:
        import google.generativeai as genai

        genai.configure(api_key=gemini_key)
        image_part = {"inline_data": {"mime_type": mime_type, "data": image_bytes}}
        content_parts = [image_part, f"Контент: изображение «{subject_or_filename}»."]
        requested_questions = min(6, num_questions + 2) if exclude_questions else num_questions
        raw = _call_gemini(gemini_key, content_parts, requested_questions, exclude_questions)
        return _parse_quiz_json(raw, exclude_questions)[:num_questions]
    except Exception as e:
        logger.warning("Gemini image quiz failed: %s", e)
        return []


def generate_quiz(
    materials: list[tuple[bytes | None, str | None, str | None, str | None]],
    subject_name: str = "материалы",
    num_questions: int = 3,
    exclude_questions: list[str] | None = None,
    *,
    gemini_proxy_key: str | None = None,
    gemini_proxy_base_url: str | None = None,
    gemini_proxy_model: str = "gemini-2.0-flash",
    deepseek_key: str | None = None,
    groq_key: str | None = None,
    gemini_key: str | None = None,
) -> list[QuizQuestion]:
    """
    materials: список (data, url, mime_type, filename).
    Сначала пробует Gemini-compatible endpoint (текст), затем DeepSeek, затем Groq, при необходимости Gemini.
    """
    if not gemini_proxy_key and not deepseek_key and not groq_key and not gemini_key:
        return []

    all_text_parts: list[str] = []
    image_parts: list[tuple[bytes, str]] = []

    for data, url, mime, fname in materials:
        if data:
            mime_l = (mime or "").lower()
            if mime_l.startswith("image/") or (fname or "").lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".gif")):
                image_parts.append((data, mime_l or "image/jpeg"))
            else:
                t = extract_text_from_material(data, None, mime, fname)
                if t:
                    all_text_parts.append(f"--- {fname or 'файл'} ---\n{t}")

    if not all_text_parts and not image_parts:
        return []

    questions: list[QuizQuestion] = []
    combined_text = "\n\n".join(all_text_parts)
    if combined_text:
        questions = generate_quiz_from_text(
            combined_text, subject_name, num_questions,
            exclude_questions=exclude_questions,
            gemini_proxy_key=gemini_proxy_key,
            gemini_proxy_base_url=gemini_proxy_base_url,
            gemini_proxy_model=gemini_proxy_model,
            deepseek_key=deepseek_key, groq_key=groq_key, gemini_key=gemini_key,
        )

    for img_data, img_mime in image_parts[:2]:
        qs = generate_quiz_from_image(
            img_data, img_mime, subject_name, min(2, num_questions - len(questions)),
            exclude_questions=[q.question for q in questions] + list(exclude_questions or []),
            gemini_key=gemini_key,
        )
        questions.extend(qs)

    return questions[:5]
